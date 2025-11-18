#!/usr/bin/env python3
"""
Customer Output Analysis Script

Analyzes the customer's CORRECT output to understand:
1. What data values are present
2. What structure the data has
3. What our extraction + rendering should produce
"""

import sys
from pathlib import Path
from docx import Document
import re

def extract_all_text_and_tables(docx_path):
    """Extract all content from a DOCX file"""

    doc = Document(docx_path)

    result = {
        "filename": Path(docx_path).name,
        "paragraphs": [],
        "tables": [],
        "all_text": [],
    }

    print(f"\n{'='*80}")
    print(f"ANALYZING: {Path(docx_path).name}")
    print(f"{'='*80}\n")

    # Extract paragraphs
    print("📄 PARAGRAPHS:\n")
    for i, para in enumerate(doc.paragraphs, 1):
        text = para.text.strip()
        if text:
            result["paragraphs"].append(text)
            result["all_text"].append(text)
            print(f"  [{i}] {text[:100]}{'...' if len(text) > 100 else ''}")

    # Extract tables
    print(f"\n📊 TABLES ({len(doc.tables)} total):\n")
    for table_idx, table in enumerate(doc.tables, 1):
        print(f"  Table {table_idx}: {len(table.rows)} rows × {len(table.columns)} cols")

        table_data = []
        for row_idx, row in enumerate(table.rows):
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_data.append(cell_text)
                if cell_text:
                    result["all_text"].append(cell_text)
            table_data.append(row_data)

            # Show first few rows
            if row_idx < 5:
                print(f"    Row {row_idx + 1}: {' | '.join([c[:30] for c in row_data])}")

        if len(table.rows) > 5:
            print(f"    ... and {len(table.rows) - 5} more rows")

        result["tables"].append(table_data)
        print()

    return result


def detect_key_fields(content):
    """Detect key COC fields in the content"""

    all_text = " ".join(content["all_text"])

    print(f"\n{'='*80}")
    print("KEY FIELDS DETECTION")
    print(f"{'='*80}\n")

    # Contract number pattern
    contract_matches = re.findall(r'(\d{3}\.\d{2}\.\d{4}\.\d{2})', all_text)
    if contract_matches:
        print(f"✅ Contract Number(s): {', '.join(set(contract_matches))}")

    # Shipment number pattern
    shipment_matches = re.findall(r'([A-Z0-9]{8,12})', all_text)
    # Filter to likely shipment numbers (starts with digit + letters)
    shipment_likely = [s for s in shipment_matches if re.match(r'\d[A-Z]{2}\d{6}', s)]
    if shipment_likely:
        print(f"✅ Shipment Number(s): {', '.join(set(shipment_likely))}")

    # Date patterns
    date_matches = re.findall(r'(\d{2}[./]\d{2}[./]\d{4})', all_text)
    if date_matches:
        print(f"✅ Date(s): {', '.join(set(date_matches))}")

    # Quantity patterns
    qty_matches = re.findall(r'(?:Quantity|Qty)[\s:]+(\d+)', all_text, re.IGNORECASE)
    if qty_matches:
        print(f"✅ Quantity: {', '.join(set(qty_matches))}")

    # Serial number count (look for patterns like "SN 1", "SN 2", etc.)
    serial_patterns = re.findall(r'(?:S/?N|Serial)\s*(?:Number)?\s*:?\s*(\d+)', all_text, re.IGNORECASE)
    if serial_patterns:
        print(f"✅ Serial References: {len(serial_patterns)} found")

    # Product codes (8 digits)
    product_codes = re.findall(r'\b(\d{11})\b', all_text)
    if product_codes:
        print(f"✅ Product Code(s): {', '.join(set(product_codes[:3]))}{'...' if len(product_codes) > 3 else ''}")

    print()


def analyze_serial_numbers(content):
    """Analyze how serial numbers are stored in the document"""

    print(f"\n{'='*80}")
    print("SERIAL NUMBER ANALYSIS")
    print(f"{'='*80}\n")

    all_text = " ".join(content["all_text"])

    # Look for serial number patterns
    # Pattern 1: "SN: 12345678"
    sn_pattern1 = re.findall(r'S/?N\s*:?\s*(\d{8,})', all_text, re.IGNORECASE)

    # Pattern 2: "Serial Number: 12345678"
    sn_pattern2 = re.findall(r'Serial\s*Number\s*:?\s*(\d{8,})', all_text, re.IGNORECASE)

    # Pattern 3: Just 8+ digit numbers (might be serials)
    potential_serials = re.findall(r'\b(\d{8,11})\b', all_text)

    total_found = len(sn_pattern1) + len(sn_pattern2)

    if total_found > 0:
        print(f"✅ Found {total_found} explicit serial number references")
        if total_found <= 10:
            for sn in (sn_pattern1 + sn_pattern2):
                print(f"  • {sn}")
        else:
            print(f"  First 5:")
            for sn in (sn_pattern1 + sn_pattern2)[:5]:
                print(f"  • {sn}")
            print(f"  ... and {total_found - 5} more")

    if potential_serials and not total_found:
        print(f"⚠️  Found {len(potential_serials)} potential serial numbers (8-11 digits)")
        print(f"  First 10: {', '.join(potential_serials[:10])}")

    # Check if serials are in a table
    if content["tables"]:
        print(f"\n📊 Checking tables for serial numbers...")
        for table_idx, table in enumerate(content["tables"], 1):
            serial_count = 0
            for row in table:
                for cell in row:
                    if re.search(r'\d{8,11}', cell):
                        serial_count += 1

            if serial_count > 5:
                print(f"  Table {table_idx}: Contains {serial_count} cells with 8+ digit numbers (likely serials)")

    print()


def main():
    """Main analysis function"""

    # Path to customer's CORRECT output
    customer_output = Path("C:/Projects/Documents/#2/COC_SV_Del165_20.03.2025.docx")

    # Path to our INCORRECT output
    our_output = Path("C:/Users/user35/Desktop/1/2/COC_SV_Del100_18.11.2025.docx")

    print("\n" + "="*80)
    print("CUSTOMER OUTPUT ANALYSIS")
    print("="*80)
    print("\nThis script analyzes the customer's CORRECT output to understand")
    print("what data should be present in our rendered document.\n")

    if not customer_output.exists():
        print(f"❌ Customer output not found: {customer_output}")
        print("\nPlease provide the correct path to the customer's output file.")
        return 1

    # Analyze customer's correct output
    print("\n" + "*"*80)
    print("CUSTOMER'S CORRECT OUTPUT")
    print("*"*80)

    customer_content = extract_all_text_and_tables(str(customer_output))
    detect_key_fields(customer_content)
    analyze_serial_numbers(customer_content)

    # Analyze our output if it exists
    if our_output.exists():
        print("\n" + "*"*80)
        print("OUR APPLICATION'S OUTPUT")
        print("*"*80)

        our_content = extract_all_text_and_tables(str(our_output))
        detect_key_fields(our_content)
        analyze_serial_numbers(our_content)

        # Compare
        print("\n" + "="*80)
        print("COMPARISON")
        print("="*80)

        customer_text_count = len(customer_content["all_text"])
        our_text_count = len(our_content["all_text"])

        print(f"\nCustomer's output: {customer_text_count} text elements")
        print(f"Our output:        {our_text_count} text elements")

        if our_text_count < customer_text_count * 0.5:
            print("\n❌ CRITICAL: Our output has significantly less data!")
            print("   This confirms the rendering is not populating fields correctly.")

    print("\n" + "="*80)
    print("\nAnalysis complete. Use this information to:")
    print("1. Verify our PDF extraction gets all these values")
    print("2. Ensure our template mapping provides these fields")
    print("3. Check that serial numbers are properly structured")
    print("="*80 + "\n")

    return 0


if __name__ == "__main__":
    sys.exit(main())
