#!/usr/bin/env python3
"""
DOCX Comparison Script
Compares two DOCX files to identify differences in content
"""

import sys
from pathlib import Path
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
import json

def extract_docx_content(docx_path):
    """Extract all content from a DOCX file"""
    doc = Document(docx_path)

    result = {
        "filename": Path(docx_path).name,
        "paragraphs": [],
        "tables": [],
        "metadata": {}
    }

    # Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:  # Only include non-empty paragraphs
            result["paragraphs"].append({
                "text": text,
                "style": para.style.name if para.style else "Normal"
            })

    # Extract tables
    for table_idx, table in enumerate(doc.tables):
        table_data = {
            "table_number": table_idx + 1,
            "rows": len(table.rows),
            "cols": len(table.columns),
            "data": []
        }

        for row_idx, row in enumerate(table.rows):
            row_data = []
            for cell_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                row_data.append(cell_text)
            table_data["data"].append(row_data)

        result["tables"].append(table_data)

    # Extract metadata
    core_props = doc.core_properties
    result["metadata"] = {
        "author": core_props.author or "",
        "title": core_props.title or "",
        "created": str(core_props.created) if core_props.created else "",
        "modified": str(core_props.modified) if core_props.modified else ""
    }

    return result


def print_section(title, char="="):
    """Print a formatted section header"""
    print(f"\n{char * 80}")
    print(f"{title:^80}")
    print(f"{char * 80}\n")


def print_docx_content(content, label):
    """Pretty print DOCX content"""
    print_section(f"{label}: {content['filename']}", "=")

    # Print metadata
    print("📋 METADATA:")
    for key, value in content["metadata"].items():
        if value:
            print(f"  {key}: {value}")

    # Print paragraphs
    print(f"\n📝 PARAGRAPHS ({len(content['paragraphs'])} total):")
    for idx, para in enumerate(content["paragraphs"], 1):
        print(f"\n  [{idx}] Style: {para['style']}")
        print(f"      Text: {para['text'][:200]}")  # Show first 200 chars
        if len(para['text']) > 200:
            print(f"      ... ({len(para['text'])} chars total)")

    # Print tables
    print(f"\n📊 TABLES ({len(content['tables'])} total):")
    for table in content["tables"]:
        print(f"\n  Table {table['table_number']}: {table['rows']} rows × {table['cols']} cols")

        # Print first 5 rows as sample
        max_rows = min(5, len(table['data']))
        for row_idx in range(max_rows):
            row = table['data'][row_idx]
            print(f"    Row {row_idx + 1}: {' | '.join([cell[:30] for cell in row])}")

        if len(table['data']) > 5:
            print(f"    ... ({len(table['data']) - 5} more rows)")


def compare_content(expected, actual):
    """Compare two DOCX contents and highlight differences"""
    print_section("COMPARISON ANALYSIS", "*")

    # Compare paragraph counts
    exp_para_count = len(expected["paragraphs"])
    act_para_count = len(actual["paragraphs"])

    print("📊 STRUCTURAL COMPARISON:")
    print(f"  Paragraphs: Expected={exp_para_count}, Actual={act_para_count}")
    if exp_para_count != act_para_count:
        print(f"  ❌ DIFFERENCE: {abs(exp_para_count - act_para_count)} paragraph(s) difference")
    else:
        print(f"  ✅ Same paragraph count")

    # Compare table counts
    exp_table_count = len(expected["tables"])
    act_table_count = len(actual["tables"])

    print(f"\n  Tables: Expected={exp_table_count}, Actual={act_table_count}")
    if exp_table_count != act_table_count:
        print(f"  ❌ DIFFERENCE: {abs(exp_table_count - act_table_count)} table(s) difference")
    else:
        print(f"  ✅ Same table count")

    # Compare table structures
    if exp_table_count > 0 and act_table_count > 0:
        print("\n📋 TABLE STRUCTURE COMPARISON:")
        for idx in range(min(exp_table_count, act_table_count)):
            exp_table = expected["tables"][idx]
            act_table = actual["tables"][idx]

            print(f"\n  Table {idx + 1}:")
            print(f"    Expected: {exp_table['rows']} rows × {exp_table['cols']} cols")
            print(f"    Actual:   {act_table['rows']} rows × {act_table['cols']} cols")

            if exp_table['rows'] != act_table['rows'] or exp_table['cols'] != act_table['cols']:
                print(f"    ❌ STRUCTURE MISMATCH")
            else:
                print(f"    ✅ Same structure")

    # Content comparison
    print("\n🔍 CONTENT ANALYSIS:")

    # Check for empty content in actual
    empty_cells = 0
    total_cells = 0

    for table in actual["tables"]:
        for row in table["data"]:
            for cell in row:
                total_cells += 1
                if not cell or cell.strip() == "":
                    empty_cells += 1

    if total_cells > 0:
        empty_percentage = (empty_cells / total_cells) * 100
        print(f"  Empty cells in ACTUAL: {empty_cells}/{total_cells} ({empty_percentage:.1f}%)")

        if empty_percentage > 50:
            print(f"  ❌ CRITICAL: More than 50% of cells are empty!")
        elif empty_percentage > 20:
            print(f"  ⚠️  WARNING: More than 20% of cells are empty")
        else:
            print(f"  ✅ Acceptable empty cell ratio")

    # Sample data comparison (first table, first 3 rows)
    if exp_table_count > 0 and act_table_count > 0:
        print("\n📋 SAMPLE DATA COMPARISON (Table 1, First 3 Rows):")
        exp_table = expected["tables"][0]
        act_table = actual["tables"][0]

        max_rows = min(3, len(exp_table['data']), len(act_table['data']))

        for row_idx in range(max_rows):
            print(f"\n  Row {row_idx + 1}:")
            exp_row = exp_table['data'][row_idx]
            act_row = act_table['data'][row_idx]

            max_cols = min(len(exp_row), len(act_row))
            for col_idx in range(max_cols):
                exp_cell = exp_row[col_idx][:50]
                act_cell = act_row[col_idx][:50]

                if exp_cell != act_cell:
                    print(f"    Col {col_idx + 1}: DIFFERENT")
                    print(f"      Expected: '{exp_cell}'")
                    print(f"      Actual:   '{act_cell}'")
                else:
                    print(f"    Col {col_idx + 1}: ✅ Same")


def find_missing_fields(expected, actual):
    """Identify specific fields that are missing in actual"""
    print_section("MISSING FIELDS ANALYSIS", "#")

    # Extract all non-empty text from expected
    exp_texts = set()
    for para in expected["paragraphs"]:
        if para["text"]:
            exp_texts.add(para["text"])

    for table in expected["tables"]:
        for row in table["data"]:
            for cell in row:
                if cell and len(cell.strip()) > 2:  # Ignore very short strings
                    exp_texts.add(cell.strip())

    # Extract all non-empty text from actual
    act_texts = set()
    for para in actual["paragraphs"]:
        if para["text"]:
            act_texts.add(para["text"])

    for table in actual["tables"]:
        for row in table["data"]:
            for cell in row:
                if cell and len(cell.strip()) > 2:
                    act_texts.add(cell.strip())

    # Find texts in expected but not in actual
    missing = exp_texts - act_texts

    if missing:
        print(f"❌ FOUND {len(missing)} UNIQUE TEXT STRINGS IN EXPECTED BUT NOT IN ACTUAL:\n")
        for idx, text in enumerate(sorted(missing), 1):
            if len(text) > 100:
                print(f"  {idx}. {text[:100]}...")
            else:
                print(f"  {idx}. {text}")
    else:
        print("✅ No obvious missing text strings detected")

    # Look for specific patterns
    print("\n🔍 CHECKING SPECIFIC FIELD PATTERNS:")

    patterns = {
        "Contract Number": r"\d+\.\d+\.\d+\.\d+",
        "Shipment Number": r"\d{1,2}[A-Z]{2}\d{6}",
        "Serial Numbers (NL)": r"NL\d{5}",
        "COC Number": r"COC\d{6}",
        "Date": r"\d{2}[./]\d{2}[./]\d{4}",
    }

    import re

    for pattern_name, pattern in patterns.items():
        exp_matches = set()
        act_matches = set()

        # Search in expected
        for text in exp_texts:
            matches = re.findall(pattern, str(text))
            exp_matches.update(matches)

        # Search in actual
        for text in act_texts:
            matches = re.findall(pattern, str(text))
            act_matches.update(matches)

        print(f"\n  {pattern_name}:")
        print(f"    Expected: {len(exp_matches)} found - {list(exp_matches)[:5]}")
        print(f"    Actual:   {len(act_matches)} found - {list(act_matches)[:5]}")

        if len(exp_matches) > len(act_matches):
            print(f"    ❌ MISSING: {len(exp_matches) - len(act_matches)} items")
        elif len(exp_matches) == len(act_matches):
            print(f"    ✅ Same count")


def main():
    """Main comparison function"""
    # File paths (Windows paths)
    expected_path = r"C:\Projects\Documents\#2\COC_SV_Del165_20.03.2025.docx"
    actual_path = r"C:\Users\user35\Desktop\1\2\COC_SV_Del100_18.11.2025.docx"

    print("=" * 80)
    print("DOCX COMPARISON TOOL".center(80))
    print("Comparing Expected vs Actual Output".center(80))
    print("=" * 80)

    # Check if files exist
    if not Path(expected_path).exists():
        print(f"\n❌ ERROR: Expected file not found: {expected_path}")
        print("   Please verify the path and try again.")
        return 1

    if not Path(actual_path).exists():
        print(f"\n❌ ERROR: Actual file not found: {actual_path}")
        print("   Please verify the path and try again.")
        return 1

    print(f"\n📂 Reading files...")
    print(f"   Expected: {expected_path}")
    print(f"   Actual:   {actual_path}")

    try:
        # Extract content
        expected = extract_docx_content(expected_path)
        actual = extract_docx_content(actual_path)

        # Print detailed content
        print_docx_content(expected, "EXPECTED OUTPUT")
        print_docx_content(actual, "ACTUAL OUTPUT")

        # Compare
        compare_content(expected, actual)

        # Find missing fields
        find_missing_fields(expected, actual)

        # Save to JSON for further analysis
        output_file = "docx_comparison_results.json"
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump({
                "expected": expected,
                "actual": actual
            }, f, indent=2, ensure_ascii=False)

        print_section("ANALYSIS COMPLETE", "=")
        print(f"✅ Detailed results saved to: {output_file}")
        print("\nNext steps:")
        print("  1. Review the comparison output above")
        print("  2. Check the JSON file for detailed data")
        print("  3. Focus on fixing empty cells and missing fields")

        return 0

    except Exception as e:
        print(f"\n❌ ERROR: {str(e)}")
        import traceback
        traceback.print_exc()
        return 1


if __name__ == "__main__":
    sys.exit(main())
