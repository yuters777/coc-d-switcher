# backend/scripts/pdf_to_template.py
from pdf2docx import Converter
from docx import Document
from pathlib import Path

def convert_pdf_to_template():
    # Get script directory and locate PDF
    script_dir = Path(__file__).parent
    pdf_path = script_dir / "COC_SV_Del165_20.03.2025.pdf"
    temp_docx = script_dir / "temp_converted.docx"
    
    if not pdf_path.exists():
        print(f"ERROR: PDF not found at {pdf_path}")
        return
    
    print(f"Converting: {pdf_path}")
    
    cv = Converter(str(pdf_path))
    cv.convert(str(temp_docx))
    cv.close()
    
    print("PDF converted, adding placeholders...")
    
    doc = Document(str(temp_docx))
    
    replacements = {
        'Supplier Serial No:': 'Supplier Serial No: {{ supplier_serial_no }}',
        'Contract Number:': 'Contract Number: {{ contract_number }}',
        'Acquirer (include Name, Address, email etc.)': 'Acquirer (include Name, Address, email etc.)\n{{ acquirer }}',
        'Delivery Address:': 'Delivery Address:\n{{ delivery_address }}',
        'Partial Delivery Number:': 'Partial Delivery Number: {{ partial_delivery_number }}',
        'Contract Item': 'Contract Item\n{{ contract_item }}',
        'Product Description or Part': 'Product Description or Part\n{{ product_description }}',
        'Quantity': 'Quantity\n{{ quantity }}',
        'Packing Slip:': 'Packing Slip: {{ shipment_no }}',
        '4196 (of 8115)': '{{ undelivered_quantity }}',
        'SW Ver. # 2.2.15.45': '{{ remarks }}'
    }
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for old, new in replacements.items():
                    if old in cell.text:
                        for paragraph in cell.paragraphs:
                            if old in paragraph.text:
                                paragraph.text = paragraph.text.replace(old, new)
    
    # Insert serial numbers page
    for i, para in enumerate(doc.paragraphs):
        if "Part II" in para.text:
            para.insert_paragraph_before("\n\nSerial Numbers\nTotal: {{ serial_count }} units\n{{ serials_list }}\n")
            break
    
    # Save to backend/templates/
    output_dir = script_dir.parent / "templates"
    output_dir.mkdir(exist_ok=True)
    output_path = output_dir / "dutch_coc_template.docx"
    
    doc.save(str(output_path))
    temp_docx.unlink()
    
    print(f"Template created: {output_path}")

if __name__ == "__main__":
    convert_pdf_to_template()