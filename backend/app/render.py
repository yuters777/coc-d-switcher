from pathlib import Path
from typing import Dict, Any
import json
import os
import logging
from docx import Document
from docxtpl import DocxTemplate

logger = logging.getLogger(__name__)

TEMPLATE_PATH = os.getenv("TEMPLATE_PATH", "templates/COC_SV_Del165_20.03.2025.docx")

def render_docx(conv_json: Dict[str, Any], job_id: str) -> Path:
    """Render DOCX from conversion data"""
    try:
        out_path = Path(f"/tmp/coc-{job_id}.docx")

        # Check if template exists
        template_file = Path(TEMPLATE_PATH)

        if template_file.exists():
            # Use template if available
            try:
                doc = DocxTemplate(template_file)
                # Prepare context from conversion JSON
                context = conv_json.get("render_vars", {})
                context.update(conv_json.get("part_I", {}))
                context.update(conv_json.get("part_II", {}))
                doc.render(context)
                doc.save(str(out_path))
                logger.info(f"Rendered DOCX using template: {template_file}")
            except Exception as template_error:
                logger.warning(f"Failed to use template, creating basic DOCX: {str(template_error)}")
                # Fallback to basic DOCX
                _create_basic_docx(conv_json, out_path)
        else:
            # Create a basic DOCX document with the data
            logger.warning(f"Template not found at {template_file}, creating basic DOCX")
            _create_basic_docx(conv_json, out_path)

        return out_path
    except Exception as e:
        logger.error(f"Error rendering DOCX for job {job_id}: {str(e)}")
        raise

def _create_basic_docx(conv_json: Dict[str, Any], out_path: Path) -> None:
    """Create a basic DOCX document with conversion data"""
    try:
        doc = Document()
        doc.add_heading('Certificate of Conformity', 0)

        # Add Part I data
        if conv_json.get("part_I"):
            doc.add_heading('Part I', level=1)
            for key, value in conv_json["part_I"].items():
                doc.add_paragraph(f"{key}: {value}")

        # Add Part II data
        if conv_json.get("part_II"):
            doc.add_heading('Part II', level=1)
            for key, value in conv_json["part_II"].items():
                doc.add_paragraph(f"{key}: {value}")

        # Add metadata comment
        doc.add_paragraph()
        doc.add_paragraph(f"Generated document - Template data: {json.dumps(conv_json.get('render_vars', {}))}")

        doc.save(str(out_path))
        logger.info(f"Created basic DOCX at {out_path}")
    except Exception as e:
        logger.error(f"Error creating basic DOCX: {str(e)}")
        raise

def convert_to_pdf(docx_path: Path) -> Path:
    """Convert DOCX to PDF"""
    try:
        pdf_path = docx_path.with_suffix(".pdf")

        # Check if file exists
        if not docx_path.exists():
            raise FileNotFoundError(f"DOCX file not found: {docx_path}")

        # In production, use LibreOffice headless conversion
        # For now, create a placeholder PDF
        # TODO: Implement actual PDF conversion using LibreOffice or other tool
        # Example: subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', str(docx_path)])

        pdf_path.write_text(f"PDF version of {docx_path.name}\nNote: This is a placeholder. Implement LibreOffice conversion for production.")
        logger.warning(f"Created placeholder PDF at {pdf_path}. Implement actual PDF conversion.")

        return pdf_path
    except Exception as e:
        logger.error(f"Error converting to PDF: {str(e)}")
        raise
