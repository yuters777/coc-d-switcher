#!/usr/bin/env python3
"""
Comprehensive template diagnostics to find why variables aren't being replaced.
"""

from pathlib import Path
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.table import _Cell, Table
import re
from lxml import etree

def analyze_template(template_path: Path):
    """Analyze template to find where variables are and why they're not being replaced."""

    print("="*70)
    print(f"TEMPLATE DIAGNOSTICS: {template_path.name}")
    print("="*70)

    doc = Document(template_path)

    issues_found = []
    variables_found = []

    # 1. Check for text boxes
    print("\n1. Checking for text boxes...")
    text_boxes = []
    for element in doc.element.body.iter():
        # Check for textboxes (w:txbxContent)
        if element.tag.endswith('txbxContent'):
            text_boxes.append(element)

    if text_boxes:
        print(f"   ❌ FOUND {len(text_boxes)} TEXT BOX(ES)")
        issues_found.append(f"{len(text_boxes)} text box(es) found - docxtpl CANNOT process these")

        for idx, tb in enumerate(text_boxes):
            text = ''.join(tb.itertext())
            if '{{' in text:
                print(f"      Text box {idx+1} contains: {text[:100]}")
                variables_found.append(('text_box', text))
    else:
        print("   ✅ No text boxes found")

    # 2. Check for content controls
    print("\n2. Checking for content controls...")
    content_controls = []
    for element in doc.element.body.iter():
        if element.tag.endswith('sdt'):  # Structured Document Tag (content control)
            content_controls.append(element)

    if content_controls:
        print(f"   ⚠️  FOUND {len(content_controls)} CONTENT CONTROL(S)")
        issues_found.append(f"{len(content_controls)} content control(s) found - may cause issues")

        for idx, cc in enumerate(content_controls):
            text = ''.join(cc.itertext())
            if '{{' in text:
                print(f"      Control {idx+1} contains: {text[:100]}")
                variables_found.append(('content_control', text))
    else:
        print("   ✅ No content controls found")

    # 3. Check paragraphs for variables
    print("\n3. Checking paragraphs for {{ }} variables...")
    para_vars = 0
    for para in doc.paragraphs:
        if '{{' in para.text and '}}' in para.text:
            vars_in_para = re.findall(r'\{\{[^}]+\}\}', para.text)
            if vars_in_para:
                para_vars += len(vars_in_para)
                print(f"   ✅ Found in paragraph: {vars_in_para}")
                variables_found.append(('paragraph', para.text[:100]))

    print(f"   Total variables in paragraphs: {para_vars}")

    # 4. Check tables
    print("\n4. Checking tables for {{ }} variables...")
    table_vars = 0
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                cell_text = cell.text
                if '{{' in cell_text and '}}' in cell_text:
                    vars_in_cell = re.findall(r'\{\{[^}]+\}\}', cell_text)
                    if vars_in_cell:
                        table_vars += len(vars_in_cell)
                        print(f"   ✅ Table {table_idx}, Row {row_idx}, Cell {cell_idx}: {vars_in_cell}")
                        variables_found.append(('table_cell', cell_text[:100]))

                        # Check if the cell contains nested text boxes
                        for element in cell._element.iter():
                            if element.tag.endswith('txbxContent'):
                                print(f"      ❌ WARNING: Variable is inside a TEXT BOX within this cell!")
                                issues_found.append(f"Table {table_idx} cell has text box containing variables")

    print(f"   Total variables in tables: {table_vars}")

    # 5. Check for split runs
    print("\n5. Checking for split {{ }} tags across runs...")
    split_issues = 0
    for para in doc.paragraphs:
        if '{{' in para.text:
            # Check if {{ and }} are in different runs
            run_texts = [run.text for run in para.runs]
            combined = ''.join(run_texts)

            # Find complete tags in combined text
            complete_tags = re.findall(r'\{\{[^}]+\}\}', combined)

            # Check if any tag is split
            for tag in complete_tags:
                tag_found_complete = False
                for run_text in run_texts:
                    if tag in run_text:
                        tag_found_complete = True
                        break

                if not tag_found_complete:
                    split_issues += 1
                    print(f"   ⚠️  SPLIT TAG: {tag}")
                    print(f"      Runs: {run_texts}")
                    issues_found.append(f"Split tag: {tag}")

    if split_issues == 0:
        print("   ✅ No split tags found")

    # 6. Summary
    print("\n" + "="*70)
    print("SUMMARY")
    print("="*70)

    total_vars = len([v for v in variables_found if '{{' in v[1]])
    print(f"\nTotal variables found: {total_vars}")
    print(f"  - In paragraphs: {para_vars}")
    print(f"  - In table cells: {table_vars}")
    print(f"  - In text boxes: {len([v for v in variables_found if v[0] == 'text_box'])}")
    print(f"  - In content controls: {len([v for v in variables_found if v[0] == 'content_control'])}")

    if issues_found:
        print(f"\n❌ ISSUES FOUND ({len(issues_found)}):")
        for issue in issues_found:
            print(f"   - {issue}")
        print("\n⚠️  docxtpl can only process variables in:")
        print("   ✅ Regular paragraph text")
        print("   ✅ Table cell text (NOT in text boxes inside cells)")
        print("   ✅ Headers/footers")
        print("\n   It CANNOT process variables in:")
        print("   ❌ Text boxes")
        print("   ❌ Shapes")
        print("   ❌ Content controls (sometimes)")
        print("   ❌ Form fields")
    else:
        print("\n✅ No obvious issues found!")
        print("   Template should work with docxtpl.")
        print("   If variables still don't render, check:")
        print("   - Variable names match exactly (case-sensitive)")
        print("   - No extra spaces in {{ variable }}")
        print("   - Data is being passed correctly to render()")

    # 7. List all unique variable names found
    print("\n" + "="*70)
    print("VARIABLE NAMES FOUND")
    print("="*70)

    all_text = []
    for para in doc.paragraphs:
        all_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text.append(cell.text)

    all_vars = set()
    for text in all_text:
        found = re.findall(r'\{\{\s*(\w+)\s*\}\}', text)
        all_vars.update(found)

    if all_vars:
        print("\nVariable names (without {{ }}):")
        for var in sorted(all_vars):
            print(f"   - {var}")
    else:
        print("\n⚠️  NO VARIABLES FOUND IN TEMPLATE!")
        print("   This might mean:")
        print("   - Variables are in text boxes")
        print("   - Variables have unusual formatting")
        print("   - Template doesn't contain {{ }} tags")

if __name__ == "__main__":
    import sys

    # Check multiple possible locations
    possible_paths = [
        Path("templates/d0d00cd7-54a4-4925-a5bd-6965624e82b8_temp_dutch_coc_template.docx"),
        Path("templates/COC_SV_Del165_20.03.2025.docx"),
        Path("../templates/d0d00cd7-54a4-4925-a5bd-6965624e82b8_temp_dutch_coc_template.docx"),
        Path("../templates/COC_SV_Del165_20.03.2025.docx"),
    ]

    # If filename provided as argument, use that
    if len(sys.argv) > 1:
        template_path = Path(sys.argv[1])
    else:
        # Find first existing path
        template_path = None
        for path in possible_paths:
            if path.exists():
                template_path = path
                break

    if not template_path or not template_path.exists():
        print(f"ERROR: Template not found!")
        print(f"\nSearched in:")
        for path in possible_paths:
            print(f"  - {path.absolute()}")
        print(f"\nUsage: python diagnose_template.py [path/to/template.docx]")
        exit(1)

    analyze_template(template_path)
