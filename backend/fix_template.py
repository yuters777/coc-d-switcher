#!/usr/bin/env python3
"""
Fix template by merging split jinja2 tags.

This script finds all {{ variable }} tags that may be split across
multiple XML runs and merges them into single runs so docxtpl can parse them.
"""

from pathlib import Path
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn
import re

def fix_template_tags(docx_path: Path, output_path: Path):
    """
    Fix jinja2 tags in a DOCX template by merging split runs.

    Args:
        docx_path: Path to input template
        output_path: Path to save fixed template
    """
    print(f"Loading template: {docx_path}")
    doc = Document(docx_path)

    fixes_made = 0

    # Process all paragraphs
    for paragraph in doc.paragraphs:
        # Get paragraph text
        text = paragraph.text

        # Check if it contains jinja2 tags
        if '{{' in text and '}}' in text:
            # Find all complete tags
            tags = re.findall(r'\{\{[^}]+\}\}', text)
            if tags:
                print(f"Found tags in paragraph: {tags}")

                # Clear all runs
                for run in paragraph.runs:
                    run.text = ''

                # Add back the complete text as a single run
                # This merges all split tags
                if paragraph.runs:
                    paragraph.runs[0].text = text
                else:
                    paragraph.add_run(text)

                fixes_made += len(tags)

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text

                    if '{{' in text and '}}' in text:
                        tags = re.findall(r'\{\{[^}]+\}\}', text)
                        if tags:
                            print(f"Found tags in table cell: {tags}")

                            # Clear all runs
                            for run in paragraph.runs:
                                run.text = ''

                            # Add back as single run
                            if paragraph.runs:
                                paragraph.runs[0].text = text
                            else:
                                paragraph.add_run(text)

                            fixes_made += len(tags)

    print(f"\nFixed {fixes_made} template tags")
    print(f"Saving to: {output_path}")
    doc.save(output_path)
    print("Done!")

if __name__ == "__main__":
    template_path = Path("templates/COC_SV_Del165_20.03.2025.docx")
    output_path = Path("templates/COC_SV_Del165_20.03.2025_FIXED.docx")

    if not template_path.exists():
        print(f"ERROR: Template not found at {template_path}")
        exit(1)

    fix_template_tags(template_path, output_path)

    print("\n" + "="*60)
    print("Template fixed!")
    print("="*60)
    print(f"\nTo use the fixed template:")
    print(f"1. Backup the original: mv {template_path} {template_path}.backup")
    print(f"2. Use the fixed version: mv {output_path} {template_path}")
    print(f"3. Or upload the fixed version via Settings UI")
